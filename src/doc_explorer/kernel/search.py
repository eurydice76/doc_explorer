import os
import re
import subprocess
import tempfile

import docx
import PyPDF2

import logging

logger = logging.getLogger("PyPDF2")
logger.setLevel(logging.ERROR)


def is_binary(filename):
    is_bin = False
    chunksize = 1024
    fin = open(filename, 'rb')
    while 1:
        chunk = fin.read(chunksize)
        if b"\0" in chunk:
            is_bin = True
        if len(chunk) < chunksize:
            break
    fin.close()
    return is_bin


def doc_search(search_text, file, ignore_case: int = re.IGNORECASE) -> bool:
    """Returns whether a doc file contains a given search text.

    Args:
        search_text: the text to search
        file: the doc file to search the text in
        ignore_case: whether the case should be ignored

    Returns:
        whether the doc file contains a given text
    """
    try:
        doc = docx.Document(file)
    except ValueError:
        tempdir = tempfile.mkdtemp()
        subprocess.call(['soffice', '--headless', '--convert-to', 'docx', file, "--outdir", tempdir])
        basename = os.path.splitext(os.path.basename(file))[0]
        temp_file = os.path.join(tempdir, basename+".docx")
        doc = docx.Document(temp_file)
    except:
        return False

    for para in doc.paragraphs:
        if re.search(search_text, para.text, ignore_case):
            return True
    return False


def pdf_search(search_text, file, ignore_case: int = re.IGNORECASE) -> bool:
    """Returns whether a pdf file contains a given search text.

    Args:
        search_text: the text to search
        file: the pdf file to search the text in
        ignore_case: whether the case should be ignored

    Returns:
        whether the pdf file contains a given text
    """
    with open(file, 'rb') as fin:
        pdf_reader = PyPDF2.PdfReader(fin)
        for page in pdf_reader.pages:
            if re.search(search_text, page.extract_text(), ignore_case):
                return True
    return False


def text_search(search_text, file, ignore_case: int = re.IGNORECASE) -> bool:
    """Returns whether a text file contains a given search text.

    Args:
        search_text: the text to search
        file: the text file to search the text in
        ignore_case: whether the case should be ignored

    Returns:
        whether the text file contains a given text
    """
    if is_binary(file):
        return False

    with open(file, 'r') as fin:
        return True if re.search(search_text, fin.read(), ignore_case) else False


search_engines = {
    ".doc": doc_search,
    ".docx":  doc_search,
    ".pdf": pdf_search,
    ".dat": text_search,
    ".txt": text_search
}
